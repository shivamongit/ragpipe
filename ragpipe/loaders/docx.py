"""DOCX file loader."""

from __future__ import annotations

from pathlib import Path

from ragpipe.core import Document


class DocxLoader:
    """Load DOCX files into Document objects.

    Extracts text from paragraphs and tables.
    """

    def load(self, path: str | Path) -> Document:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("Install python-docx: pip install 'ragpipe[docx]'")

        p = Path(path)
        doc = DocxDocument(str(p))

        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return Document(
            content="\n\n".join(parts),
            metadata={
                "source": str(p),
                "filename": p.name,
                "type": ".docx",
                "paragraph_count": len(doc.paragraphs),
            },
        )

    def load_many(self, paths: list[str | Path]) -> list[Document]:
        return [self.load(p) for p in paths]
