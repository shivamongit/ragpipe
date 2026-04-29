"""FastAPI REST API server for ragpipe pipelines.

Endpoints:
    GET  /health                — liveness check
    GET  /providers             — list LLM providers + availability
    GET  /models                — list all available models
    GET  /stats                 — document/chunk counts
    POST /ingest                — ingest JSON documents
    POST /upload                — upload files (PDF/DOCX/TXT/MD)
    POST /query                 — RAG query (optional model/provider override)
    WS   /query/stream          — streaming RAG query
    DELETE /index               — clear all documents and chunks
    POST /evaluate              — evaluate pipeline on labeled query
    GET  /conversations         — list saved chats
    POST /conversations         — create new chat
    GET  /conversations/{id}    — get chat with all messages
    PATCH /conversations/{id}   — rename chat
    DELETE /conversations/{id}  — delete chat
"""

import io
import os
import time
from pathlib import Path
from typing import Any, Optional, Union

from ragpipe.core import Document, Pipeline


def create_app(
    pipeline: Union[Pipeline, None] = None,
    api_key: Union[str, None] = None,
    db_path: str = "ragpipe.db",
):
    """Create a FastAPI app wrapping a ragpipe Pipeline.

    Args:
        pipeline: Pre-configured Pipeline instance.
        api_key: Optional API key for authentication via X-API-Key header.
        db_path: SQLite path for conversation persistence.
    """
    try:
        from fastapi import (
            FastAPI, HTTPException, WebSocket, WebSocketDisconnect,
            Depends, Header, UploadFile, File, Form,
        )
        from fastapi.middleware.cors import CORSMiddleware
        from pydantic import BaseModel
    except ImportError:
        raise ImportError("Install server dependencies: pip install 'ragpipe[server]'")

    from ragpipe.generators.registry import (
        list_providers, find_model, build_generator,
    )
    from ragpipe.server.storage import ConversationStore

    app = FastAPI(
        title="ragpipe",
        description="Production RAG pipeline with multi-LLM provider support",
        version="3.1.0",
    )

    app.add_middleware(
        CORSMiddleware,
        allow_origins=["*"],
        allow_methods=["*"],
        allow_headers=["*"],
    )

    app.state.pipeline = pipeline
    app.state.store = ConversationStore(db_path=db_path)
    _api_key = api_key or os.environ.get("RAGPIPE_API_KEY")

    async def verify_api_key(x_api_key: Optional[str] = Header(None)):
        if _api_key and x_api_key != _api_key:
            raise HTTPException(status_code=401, detail="Invalid or missing API key")

    def get_pipeline() -> Pipeline:
        pipe = app.state.pipeline
        if pipe is None:
            raise HTTPException(status_code=503, detail="Pipeline not configured")
        return pipe

    # ── Pydantic models ──────────────────────────────────────────────────────

    class IngestRequest(BaseModel):
        documents: list[dict[str, Any]]

    class IngestResponse(BaseModel):
        documents: int
        chunks: int

    class QueryRequest(BaseModel):
        question: str
        top_k: Optional[int] = None
        model: Optional[str] = None  # e.g. "gpt-5-mini"
        provider: Optional[str] = None  # e.g. "openai"
        api_key_override: Optional[str] = None
        conversation_id: Optional[str] = None  # save into a conversation if provided

    class SourceResponse(BaseModel):
        text: str
        doc_id: str
        score: float
        rank: int

    class QueryResponse(BaseModel):
        answer: str
        sources: list[SourceResponse]
        model: str
        tokens_used: int
        latency_ms: float
        conversation_id: Optional[str] = None
        message_id: Optional[str] = None

    class StatsResponse(BaseModel):
        documents: int
        chunks: int

    class EvaluateRequest(BaseModel):
        question: str
        relevant_doc_ids: list[str]
        top_k: Optional[int] = None

    class EvaluateResponse(BaseModel):
        hit_rate: float
        mrr: float
        precision_at_k: float
        recall_at_k: float
        ndcg_at_k: float

    class ConversationCreateRequest(BaseModel):
        title: Optional[str] = "New chat"
        model: Optional[str] = None
        provider: Optional[str] = None

    class ConversationRenameRequest(BaseModel):
        title: str

    # ── Helpers ──────────────────────────────────────────────────────────────

    def _resolve_generator(req_provider: Optional[str], req_model: Optional[str], api_key_override: Optional[str]):
        """If query specifies a different model, build a generator for it; else use pipeline's default."""
        pipe = get_pipeline()
        if not req_model and not req_provider:
            return pipe.generator
        # Determine provider: explicit, or look up from model id
        provider = req_provider
        if not provider and req_model:
            found = find_model(req_model)
            if found:
                provider = found[0].id
        if not provider or not req_model:
            raise HTTPException(status_code=400, detail="Both provider and model required for override")
        try:
            return build_generator(provider, req_model, api_key=api_key_override)
        except (ImportError, ValueError) as e:
            raise HTTPException(status_code=400, detail=str(e))

    # ── Endpoints ────────────────────────────────────────────────────────────

    @app.get("/health")
    async def health():
        return {"status": "ok", "version": "3.1.0"}

    @app.get("/providers", dependencies=[Depends(verify_api_key)])
    async def get_providers():
        provs = list_providers()
        return {
            "providers": [
                {
                    "id": p.id,
                    "name": p.name,
                    "requires_api_key": p.requires_api_key,
                    "api_key_env_var": p.api_key_env_var,
                    "docs_url": p.docs_url,
                    "available": p.available,
                    "model_count": len(p.models),
                }
                for p in provs
            ]
        }

    @app.get("/models", dependencies=[Depends(verify_api_key)])
    async def get_models(provider: Optional[str] = None, available_only: bool = False):
        provs = list_providers()
        all_models = []
        for p in provs:
            if provider and p.id != provider:
                continue
            if available_only and not p.available:
                continue
            for m in p.models:
                all_models.append({
                    "id": m.id,
                    "name": m.name,
                    "provider": m.provider,
                    "provider_name": p.name,
                    "context_window": m.context_window,
                    "input_cost_per_m": m.input_cost_per_m,
                    "output_cost_per_m": m.output_cost_per_m,
                    "streaming": m.streaming,
                    "description": m.description,
                    "tags": m.tags,
                    "available": p.available,
                })
        return {"models": all_models, "total": len(all_models)}

    @app.post("/ingest", response_model=IngestResponse, dependencies=[Depends(verify_api_key)])
    async def ingest(req: IngestRequest):
        pipe = get_pipeline()
        docs = [
            Document(
                content=d.get("content", ""),
                metadata=d.get("metadata", {}),
                doc_id=d.get("doc_id", ""),
            )
            for d in req.documents
        ]
        stats = await pipe.aingest(docs)
        return IngestResponse(**stats)

    @app.post("/upload", dependencies=[Depends(verify_api_key)])
    async def upload(files: list[UploadFile] = File(...)):
        """Upload one or more files (PDF, DOCX, TXT, MD). Returns ingest stats."""
        pipe = get_pipeline()
        docs: list[Document] = []
        for uf in files:
            data = await uf.read()
            suffix = Path(uf.filename or "file.txt").suffix.lower()
            content = _extract_text_from_bytes(data, suffix, uf.filename or "")
            if not content.strip():
                continue
            docs.append(Document(
                content=content,
                metadata={"filename": uf.filename, "type": suffix, "size": len(data)},
            ))
        if not docs:
            raise HTTPException(status_code=400, detail="No readable content in uploaded files")
        stats = await pipe.aingest(docs)
        return {"documents": stats["documents"], "chunks": stats["chunks"], "files": [uf.filename for uf in files]}

    @app.post("/query", response_model=QueryResponse, dependencies=[Depends(verify_api_key)])
    async def query(req: QueryRequest):
        pipe = get_pipeline()
        generator = _resolve_generator(req.provider, req.model, req.api_key_override)

        # Run pipeline retrieve + this generator
        t0 = time.perf_counter()
        results = await pipe.aretrieve(req.question, top_k=req.top_k)
        gen_out = await generator.agenerate(req.question, results)
        latency_ms = round((time.perf_counter() - t0) * 1000, 2)

        sources = [
            SourceResponse(
                text=s.chunk.text[:500],
                doc_id=s.chunk.doc_id,
                score=round(s.score, 4),
                rank=s.rank,
            )
            for s in results
        ]

        # Persist into conversation if requested
        conv_id = req.conversation_id
        msg_id = None
        if conv_id:
            store: ConversationStore = app.state.store
            if not store.get_conversation(conv_id):
                raise HTTPException(status_code=404, detail="Conversation not found")
            # Save user msg + assistant msg
            store.add_message(conv_id, role="user", content=req.question)
            saved = store.add_message(
                conv_id, role="assistant", content=gen_out.answer,
                sources=[s.dict() for s in sources],
                model=gen_out.model, tokens_used=gen_out.tokens_used, latency_ms=latency_ms,
            )
            msg_id = saved["id"]

        return QueryResponse(
            answer=gen_out.answer,
            sources=sources,
            model=gen_out.model,
            tokens_used=gen_out.tokens_used,
            latency_ms=latency_ms,
            conversation_id=conv_id,
            message_id=msg_id,
        )

    @app.websocket("/query/stream")
    async def query_stream(ws: WebSocket):
        await ws.accept()

        if _api_key:
            try:
                auth = await ws.receive_json()
                if auth.get("api_key") != _api_key:
                    await ws.close(code=4001, reason="Invalid API key")
                    return
            except Exception:
                await ws.close(code=4001, reason="Auth required")
                return

        try:
            while True:
                data = await ws.receive_json()
                question = data.get("question", "")
                top_k = data.get("top_k")
                req_model = data.get("model")
                req_provider = data.get("provider")
                api_key_override = data.get("api_key_override")
                conv_id = data.get("conversation_id")

                pipe = get_pipeline()
                try:
                    generator = _resolve_generator(req_provider, req_model, api_key_override)
                except HTTPException as e:
                    await ws.send_json({"type": "error", "message": e.detail})
                    continue

                t0 = time.perf_counter()
                results = await pipe.aretrieve(question, top_k=top_k)
                # Send sources first
                await ws.send_json({
                    "type": "sources",
                    "sources": [
                        {
                            "text": s.chunk.text[:500],
                            "doc_id": s.chunk.doc_id,
                            "score": round(s.score, 4),
                            "rank": s.rank,
                        }
                        for s in results
                    ],
                })

                full_answer = []
                async for token in generator.astream(question, results):
                    full_answer.append(token)
                    await ws.send_json({"type": "token", "content": token})

                latency_ms = round((time.perf_counter() - t0) * 1000, 2)
                answer_text = "".join(full_answer)

                # Persist if requested
                msg_id = None
                if conv_id:
                    store: ConversationStore = app.state.store
                    if store.get_conversation(conv_id):
                        store.add_message(conv_id, role="user", content=question)
                        saved = store.add_message(
                            conv_id, role="assistant", content=answer_text,
                            sources=[
                                {
                                    "text": s.chunk.text[:500],
                                    "doc_id": s.chunk.doc_id,
                                    "score": round(s.score, 4),
                                    "rank": s.rank,
                                }
                                for s in results
                            ],
                            model=getattr(generator, "model", "unknown"),
                            latency_ms=latency_ms,
                        )
                        msg_id = saved["id"]

                await ws.send_json({
                    "type": "done",
                    "model": getattr(generator, "model", "unknown"),
                    "latency_ms": latency_ms,
                    "conversation_id": conv_id,
                    "message_id": msg_id,
                })
        except WebSocketDisconnect:
            pass

    @app.get("/stats", response_model=StatsResponse, dependencies=[Depends(verify_api_key)])
    async def stats():
        pipe = get_pipeline()
        return StatsResponse(documents=pipe.document_count, chunks=pipe.chunk_count)

    @app.delete("/index", dependencies=[Depends(verify_api_key)])
    async def delete_index():
        pipe = get_pipeline()
        pipe._documents.clear()
        # Reset retriever's internal state if possible
        if hasattr(pipe.retriever, "reset"):
            pipe.retriever.reset()
        elif hasattr(pipe.retriever, "_chunks"):
            pipe.retriever._chunks = []
            if hasattr(pipe.retriever, "_embeddings"):
                pipe.retriever._embeddings = None
        if hasattr(app.state, "kg_cache"):
            app.state.kg_cache = {"signature": None, "graph": None}
        return {"status": "ok", "message": "Index cleared"}

    @app.post("/evaluate", response_model=EvaluateResponse, dependencies=[Depends(verify_api_key)])
    async def evaluate(req: EvaluateRequest):
        from ragpipe.evaluation import hit_rate, mrr, precision_at_k, recall_at_k, ndcg_at_k

        pipe = get_pipeline()
        results = await pipe.aretrieve(req.question, top_k=req.top_k)
        relevant = set(req.relevant_doc_ids)
        k = req.top_k or pipe.top_k

        return EvaluateResponse(
            hit_rate=hit_rate(results, relevant),
            mrr=mrr(results, relevant),
            precision_at_k=precision_at_k(results, relevant, k=k),
            recall_at_k=recall_at_k(results, relevant, k=k),
            ndcg_at_k=ndcg_at_k(results, relevant, k=k),
        )

    # ── Conversation endpoints ───────────────────────────────────────────────

    @app.get("/conversations", dependencies=[Depends(verify_api_key)])
    async def list_conversations():
        return {"conversations": app.state.store.list_conversations()}

    @app.post("/conversations", dependencies=[Depends(verify_api_key)])
    async def create_conversation(req: ConversationCreateRequest):
        return app.state.store.create_conversation(
            title=req.title or "New chat",
            model=req.model,
            provider=req.provider,
        )

    @app.get("/conversations/{cid}", dependencies=[Depends(verify_api_key)])
    async def get_conversation(cid: str):
        conv = app.state.store.get_conversation(cid)
        if not conv:
            raise HTTPException(status_code=404, detail="Conversation not found")
        return conv

    @app.patch("/conversations/{cid}", dependencies=[Depends(verify_api_key)])
    async def rename_conversation(cid: str, req: ConversationRenameRequest):
        ok = app.state.store.update_conversation_title(cid, req.title)
        if not ok:
            raise HTTPException(status_code=404, detail="Conversation not found")
        return {"status": "ok"}

    @app.delete("/conversations/{cid}", dependencies=[Depends(verify_api_key)])
    async def delete_conversation(cid: str):
        ok = app.state.store.delete_conversation(cid)
        if not ok:
            raise HTTPException(status_code=404, detail="Conversation not found")
        return {"status": "ok"}

    # ── Knowledge Graph ──────────────────────────────────────────────────────

    app.state.kg_cache = {"signature": None, "graph": None}

    @app.get("/graph", dependencies=[Depends(verify_api_key)])
    async def get_graph(max_entities: int = 60, max_triples: int = 200):
        """Build or return a cached knowledge graph from indexed documents.

        Uses heuristic (regex) entity/relation extraction so it works without
        an LLM. The result is cached and only rebuilt when the document set
        changes.
        """
        from ragpipe.graph.knowledge_graph import KnowledgeGraph

        pipe = get_pipeline()
        docs = pipe._documents
        sig = (len(docs), pipe.chunk_count)

        cache = app.state.kg_cache
        if cache["signature"] != sig or cache["graph"] is None:
            kg = KnowledgeGraph()  # heuristic extraction (no LLM)
            for i, d in enumerate(docs):
                source = (
                    d.metadata.get("filename")
                    or d.metadata.get("source")
                    or f"doc-{i}"
                )
                try:
                    kg.add_document(d.content, source=source)
                except Exception:
                    continue
            cache["signature"] = sig
            cache["graph"] = kg

        kg = cache["graph"]

        # Top entities by mentions / connectivity
        entities_sorted = sorted(
            kg.entities,
            key=lambda e: (e.mentions, len(kg._adjacency.get(e.id, {}))),
            reverse=True,
        )[:max_entities]

        keep_ids = {e.id for e in entities_sorted}
        triples = [
            t.to_dict() for t in kg.triples
            if t.subject.lower().strip() in keep_ids
            and t.object.lower().strip() in keep_ids
        ][:max_triples]

        return {
            "entity_count": kg.entity_count,
            "triple_count": kg.triple_count,
            "documents_indexed": len(docs),
            "entities": [
                {
                    "id": e.id,
                    "name": e.name,
                    "mentions": e.mentions,
                    "neighbors": len(kg._adjacency.get(e.id, {})),
                }
                for e in entities_sorted
            ],
            "triples": triples,
        }

    return app


def _extract_text_from_bytes(data: bytes, suffix: str, filename: str) -> str:
    """Extract text content from uploaded file bytes based on extension."""
    if suffix in (".txt", ".md", ".markdown", ".rst", ".log", ".csv", ".json", ""):
        try:
            return data.decode("utf-8")
        except UnicodeDecodeError:
            return data.decode("utf-8", errors="replace")
    if suffix == ".pdf":
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError("PDF support: pip install 'ragpipe[pdf]'")
        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join((p.extract_text() or "") for p in reader.pages)
    if suffix in (".docx", ".doc"):
        try:
            from docx import Document as DocxDoc
        except ImportError:
            raise ImportError("DOCX support: pip install 'ragpipe[docx]'")
        d = DocxDoc(io.BytesIO(data))
        return "\n".join(p.text for p in d.paragraphs if p.text)
    if suffix in (".html", ".htm"):
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            return data.decode("utf-8", errors="replace")
        soup = BeautifulSoup(data, "html.parser")
        return soup.get_text(separator="\n")
    # Fallback: try utf-8
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return ""
